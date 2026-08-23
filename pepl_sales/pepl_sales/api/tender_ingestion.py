from __future__ import annotations

import ipaddress
import json
import re
import socket
from datetime import datetime
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import frappe
from frappe import _
from frappe.utils import now_datetime
from pypdf import PdfReader


MAX_SOURCE_PDF_BYTES = 25 * 1024 * 1024
MAX_LINK_COUNT = 40
MAX_LINK_DOWNLOAD_BYTES = 15 * 1024 * 1024
MAX_REDIRECTS = 3

ALLOWED_DOWNLOAD_EXTENSIONS = {
    ".pdf",
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
}

ALLOWED_GEM_HOSTS = {
    "gem.gov.in",
    "mkp.gem.gov.in",
    "bidplus.gem.gov.in",
}


def _clean_text(value):
    return re.sub(r"[ \t]+", " ", (value or "")).strip()


def _normalise_multiline_text(value):
    lines = []

    for line in (value or "").splitlines():
        cleaned = _clean_text(line)

        if cleaned:
            lines.append(cleaned)

    return "\n".join(lines)


def _get_attached_file_content(file_url):
    file_name = frappe.db.get_value(
        "File",
        {"file_url": file_url},
        "name",
    )

    if not file_name:
        frappe.throw(
            _("Attached file {0} was not found.").format(file_url)
        )

    file_doc = frappe.get_doc("File", file_name)

    return file_doc.get_content()


def _extract_pdf_text(reader):
    pages = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""

        pages.append(
            {
                "page": page_number,
                "text": _normalise_multiline_text(text),
            }
        )

    return pages


def _extract_pdf_links(reader):
    links = []
    seen = set()

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        annotations = page.get("/Annots") or []

        for annotation_reference in annotations:
            try:
                annotation = annotation_reference.get_object()
            except Exception:
                continue

            action = annotation.get("/A")

            if not action:
                continue

            uri = action.get("/URI")

            if not uri:
                continue

            uri = str(uri).strip()

            if not uri or uri in seen:
                continue

            seen.add(uri)

            parsed = urlparse(uri)
            hostname = (parsed.hostname or "").lower()

            links.append(
                {
                    "page": page_number,
                    "url": uri,
                    "host": hostname,
                    "link_type": (
                        "GeM Document"
                        if _is_allowed_gem_hostname(hostname)
                        else "External Link"
                    ),
                }
            )

            if len(links) >= MAX_LINK_COUNT:
                return links

    return links


def _first_match(patterns, text, flags=re.IGNORECASE):
    for pattern in patterns:
        match = re.search(pattern, text, flags)

        if match:
            value = _clean_text(match.group(1))

            if value:
                return value

    return None


def _business_yes_no_status(value):
    """Return the client-facing four-state review value."""
    value = str(value or "").strip().lower()

    if value in {"yes", "y", "true", "1"}:
        return "Yes"

    if value in {"no", "n", "false", "0"}:
        return "No"

    return "Not Detected"


def _detect_tender_participation_type(text):
    """Detect Tender type only from explicit wording.

    Do not infer Tender type merely because the document is publicly
    available on GeM.
    """
    value = text or ""

    if re.search(
        r"\bsource\s+development\b",
        value,
        re.IGNORECASE,
    ):
        return "Source Development"

    if re.search(
        r"\blimited\s+(?:tender|bid)\b",
        value,
        re.IGNORECASE,
    ):
        return "Limited"

    if re.search(
        r"\bopen\s+(?:tender|bid)\b",
        value,
        re.IGNORECASE,
    ):
        return "Open"

    return "Not Detected"



def _normalise_business_text(value):
    """Collapse PDF line breaks for label-aware Tender extraction."""
    return re.sub(
        r"\s+",
        " ",
        value or "",
    ).strip()


def _clean_business_value(value):
    value = _normalise_business_text(
        value
    )

    return value.strip(
        " \t\r\n:-/"
    )


def _unique_business_values(values):
    result = []

    for value in values:
        value = _clean_business_value(
            value
        )

        if not value:
            continue

        key = value.lower()

        if any(
            existing.lower() == key
            for existing in result
        ):
            continue

        result.append(value)

    return result


def _extract_factory_unit(text):
    """
    Extract a real factory/unit name rather than Ministry/Department
    continuation text.
    """
    value = _normalise_business_text(
        text
    )

    patterns = [
        (
            r"\b("
            r"Ammunition\s+Factory\s+"
            r"[A-Za-z0-9][A-Za-z0-9 .&()/-]{1,80}?"
            r")"
            r"(?=,\s*(?:Department|Ministry|Munitions|Organisation)"
            r"|\s*\(|\s+(?:Department|Ministry)\b)"
        ),
        (
            r"\b("
            r"Ordnance\s+Factory\s+"
            r"[A-Za-z0-9][A-Za-z0-9 .&()/-]{1,80}?"
            r")"
            r"(?=,\s*(?:Department|Ministry|Organisation)"
            r"|\s*\()"
        ),
        (
            r"\b("
            r"[A-Za-z0-9][A-Za-z0-9 .&()/-]{1,80}"
            r"\s+Factory"
            r")"
            r"(?=,\s*(?:Department|Ministry|Organisation)"
            r"|\s*\()"
        ),
    ]

    for pattern in patterns:
        match = re.search(
            pattern,
            value,
            re.IGNORECASE,
        )

        if match:
            return _clean_business_value(
                match.group(1)
            )

    return None


def _extract_technical_references(text):
    """
    Extract only strong drawing/specification identifiers.

    Generic prose containing the word 'specification' is deliberately
    ignored so phrases such as 'specification and got rejected by'
    cannot become a technical reference.
    """
    value = _normalise_business_text(
        text
    )

    values = []

    patterns = [
        (
            r"\bCIA/AMN/\d+"
            r"\s+(?:PT\.?|PART)"
            r"\s*NO\.?\s*[A-Z0-9-]+"
        ),
        (
            r"\bDES\s+REF\s+NO\.?\s*"
            r"([A-Z0-9][A-Z0-9./()\- ]{2,100}?)"
            r"(?=\s+(?:GeMARPTS|SPECN|DC|Authorized|"
            r"Category|Technical|\(\s*\d+\s+pieces)|$)"
        ),
        (
            r"\bSPECN\s+NO\.?\s*"
            r"([A-Z0-9][A-Z0-9 ()/.\-]{2,80}?)"
            r"(?=\s+(?:DC\s*\d|&\s*L\s*\d|"
            r"Authorized|Technical|\(\s*\d+\s+pieces)|$)"
        ),
        r"\bDC\s*\d+-[A-Z]\b",
        r"\bL\s*\d+-[A-Z]\b",
    ]

    for pattern in patterns:
        for match in re.finditer(
            pattern,
            value,
            re.IGNORECASE,
        ):
            candidate = (
                match.group(1)
                if match.lastindex
                else match.group(0)
            )

            values.append(
                candidate
            )

    values = _unique_business_values(
        values
    )

    if not values:
        return None

    return "; ".join(
        values[:8]
    )


def _extract_explicit_delivery_period(text):
    """
    Capture only an explicit numeric delivery period/schedule.

    Generic contractual prose mentioning delivery is ignored.
    """
    value = _normalise_business_text(
        text
    )

    patterns = [
        (
            r"\bDelivery\s+Period\b"
            r"\s*(?::|-)?\s*"
            r"("
            r"\d+(?:\s*[-–]\s*\d+)?"
            r"\s*(?:day|days|week|weeks|month|months)"
            r"(?:\s+from\s+[A-Za-z0-9 /().,-]+)?"
            r")"
        ),
        (
            r"\bDelivery\s+Schedule\b"
            r"\s*(?::|-)?\s*"
            r"("
            r"\d+(?:\s*[-–]\s*\d+)?"
            r"\s*(?:day|days|week|weeks|month|months)"
            r")"
        ),
    ]

    for pattern in patterns:
        match = re.search(
            pattern,
            value,
            re.IGNORECASE,
        )

        if match:
            return _clean_business_value(
                match.group(1)
            )

    return None


def _extract_documents_required_from_bidders(text):
    value = _normalise_business_text(
        text
    )

    patterns = [
        (
            r"\bDocument\s+required\s+from\s+seller\b"
            r"\s*(?::|-)?\s*"
            r"(.+?)"
            r"(?=\s+\*In\s+case|\s+Bid\s+Number\b|$)"
        ),
        (
            r"\bDocuments\s+Required\s+From\s+All\s+Bidders\b"
            r"\s*(?::|-)?\s*"
            r"(.+?)"
            r"(?=\s+\*In\s+case|\s+Bid\s+Number\b|$)"
        ),
    ]

    for pattern in patterns:
        match = re.search(
            pattern,
            value,
            re.IGNORECASE,
        )

        if match:
            candidate = _clean_business_value(
                match.group(1)
            )

            if candidate:
                return candidate

    return None


def _extract_inspection_required(text):
    value = _normalise_business_text(
        text
    )

    match = re.search(
        (
            r"\bInspection\s+Required\b"
            r".{0,260}?"
            r"\b(Yes|No)\b"
        ),
        value,
        re.IGNORECASE,
    )

    if not match:
        return None

    return match.group(1).title()


def _extract_inspection_type(text):
    value = _normalise_business_text(
        text
    )

    patterns = [
        (
            r"\bType\s+Of\s+Inspection\b"
            r"\s*(?::|-)?\s*"
            r"(.+?)"
            r"(?=\s+Name\s+of\s+the\s+Empanelled|\s+Auto\s+CRAC|$)"
        ),
        (
            r"\bInspection\s+Type\b"
            r"\s*(?::|-)?\s*"
            r"(.+?)"
            r"(?=\s+Name\s+of\s+the\s+Empanelled|\s+Auto\s+CRAC|$)"
        ),
    ]

    for pattern in patterns:
        match = re.search(
            pattern,
            value,
            re.IGNORECASE,
        )

        if match:
            candidate = _clean_business_value(
                match.group(1)
            )

            if candidate:
                return candidate

    return None


def _extract_inspection_authority(text):
    value = _normalise_business_text(
        text
    )

    patterns = [
        (
            r"\bName\s+of\s+the\s+Empanelled\s+Inspection"
            r"\s+Agency\s*/?\s*Authority\b"
            r"\s*(?::|-)?\s*"
            r"(.+?)"
            r"(?=\s+Auto\s+CRAC|\s+Evaluation\s+Method|$)"
        ),
        (
            r"\bInspection\s+Agency\s+Name\b"
            r"\s*(?::|-)?\s*"
            r"(.+?)"
            r"(?=\s+Auto\s+CRAC|\s+Evaluation\s+Method|$)"
        ),
    ]

    for pattern in patterns:
        match = re.search(
            pattern,
            value,
            re.IGNORECASE,
        )

        if match:
            candidate = _clean_business_value(
                match.group(1)
            )

            if candidate:
                return candidate

    return None


def _extract_tender_business_review(text, parsed):
    """Extract conservative business-facing Tender information.

    Prefer explicit labels and technical identifiers over generic prose.
    Generic GeM policy/disclaimer text must not create a business value.
    """
    result = {}

    result["factory_name"] = (
        _extract_factory_unit(
            text
        )
    )

    result["nomenclature"] = (
        _first_match(
            [
                r"Nomenclature"
                r"\s*(?::|-)?\s*([^\n]+)",
            ],
            text,
        )
        or parsed.get("item_category")
    )

    result[
        "specification_drawing_reference"
    ] = _extract_technical_references(
        text
    )

    result["reverse_auction"] = _first_match(
        [
            r"Bid\s*to\s*RA\s*enabled"
            r"\s*(?::|-)?\s*(Yes|No)",

            r"Reverse\s*Auction"
            r"\s*(?:Applicable|Enabled)?"
            r"\s*(?::|-)?\s*(Yes|No)",
        ],
        text,
    )

    # Deliberately match only explicit Advance Sample labels.
    result["advance_sample"] = _first_match(
        [
            r"Advance\s*Sample"
            r"\s*(?:Required)?"
            r"\s*(?::|-)?\s*(Yes|No)",

            r"Advance\s*Sample\s*Requirement"
            r"\s*(?::|-)?\s*(Yes|No)",
        ],
        text,
    )

    result[
        "delivery_period_requirement"
    ] = _extract_explicit_delivery_period(
        text
    )

    result["option_clause"] = _first_match(
        [
            r"Option\s*Clause"
            r"\s*(?:Applicable)?"
            r"\s*(?::|-)?\s*(Yes|No)",
        ],
        text,
    )

    result[
        "documents_required_from_bidders"
    ] = _extract_documents_required_from_bidders(
        text
    )

    result[
        "inspection_required"
    ] = _extract_inspection_required(
        text
    )

    result[
        "inspection_type"
    ] = _extract_inspection_type(
        text
    )

    result[
        "inspection_authority"
    ] = _extract_inspection_authority(
        text
    )

    result[
        "mii_purchase_preference"
    ] = _first_match(
        [
            r"MII\s*Purchase\s*Preference"
            r"\s*(?::|-)?\s*(Yes|No)",

            r"Make\s*in\s*India\s*Purchase\s*Preference"
            r"\s*(?::|-)?\s*(Yes|No)",
        ],
        text,
    )

    result[
        "mse_purchase_preference"
    ] = _first_match(
        [
            r"MSE\s*Purchase\s*Preference"
            r"\s*(?::|-)?\s*(Yes|No)",

            r"Purchase\s*Preference\s*to\s*MSE"
            r"\s*(?::|-)?\s*(Yes|No)",
        ],
        text,
    )

    result[
        "tender_participation_type"
    ] = _detect_tender_participation_type(
        text
    )

    return result


def _load_tender_extraction_payload(tender):
    raw = getattr(
        tender,
        "tender_extraction_json",
        None,
    )

    if not raw:
        return {}

    try:
        payload = json.loads(raw)

        if isinstance(payload, dict):
            return payload

    except Exception:
        pass

    return {}


def _business_field_is_safe_to_refresh(
    fieldname,
    current,
    previous_auto_value=None,
):
    """
    Refresh only empty/not-detected values or recognisable stale automatic
    values produced by the previous extraction rules.
    """
    if current in {
        None,
        "",
        "Not Detected",
        "Review Required",
    }:
        return True

    if (
        previous_auto_value not in {
            None,
            "",
        }
        and str(current).strip()
        == str(previous_auto_value).strip()
    ):
        return True

    value = str(
        current
    ).strip().lower()

    stale_patterns = {
        "factory_name": [
            "department of defence production",
            "ministry of",
        ],
        "specification_drawing_reference": [
            "got rejected",
            "rejected by",
        ],
        "delivery_period_requirement": [
            "otherwise becomes apparent",
            "inability otherwise",
        ],
        "inspection_authority": [
            "/ agencies",
        ],
    }

    return any(
        token in value
        for token in stale_patterns.get(
            fieldname,
            []
        )
    )


def _apply_business_review_enrichment(
    tender,
    extracted,
    previous_parsed=None,
):
    """
    Enrich business-review fields without overwriting an intentional,
    sensible user-reviewed value.
    """
    previous_parsed = (
        previous_parsed
        or {}
    )

    # Clean up recognisable values generated by the retired broad
    # extraction rules. These are safe to clear because the corrected
    # parser intentionally produced no valid replacement.
    stale_clear_rules = {
        "delivery_period_requirement": [
            "otherwise becomes apparent",
            "inability otherwise",
        ],
        "specification_drawing_reference": [
            "got rejected by",
            "rejected by",
        ],
        "inspection_authority": [
            "/ agencies",
        ],
    }

    for fieldname, tokens in stale_clear_rules.items():
        current = getattr(
            tender,
            fieldname,
            None,
        )

        if not current:
            continue

        value = str(
            current
        ).strip().lower()

        if any(
            token in value
            for token in tokens
        ):
            setattr(
                tender,
                fieldname,
                None,
            )

    mapping = {
        "factory_name":
            "factory_name",

        # Nomenclature is owned by the primary Tender extraction.
        # Large ATC tables commonly contain headings such as
        # "Nomenclature of Item" which can otherwise produce fragments
        # like "of". Supporting documents must not replace a valid
        # primary nomenclature.
        "specification_drawing_reference":
            "specification_drawing_reference",

        "delivery_period_requirement":
            "delivery_period_requirement",

        "documents_required_from_bidders":
            "documents_required_from_bidders",

        "inspection_type":
            "inspection_type",

        "inspection_authority":
            "inspection_authority",

        "tender_participation_type":
            "tender_participation_type",
    }

    for source_key, fieldname in mapping.items():
        value = extracted.get(
            source_key
        )

        if value in {
            None,
            "",
            "Not Detected",
        }:
            continue

        current = getattr(
            tender,
            fieldname,
            None,
        )

        previous = previous_parsed.get(
            source_key
        )

        if _business_field_is_safe_to_refresh(
            fieldname,
            current,
            previous,
        ):
            setattr(
                tender,
                fieldname,
                value,
            )

    status_mapping = {
        "reverse_auction":
            "reverse_auction_status",

        "advance_sample":
            "advance_sample_status",

        "option_clause":
            "option_clause_status",

        "inspection_required":
            "inspection_status",

        "mii_purchase_preference":
            "mii_purchase_preference_status",

        "mse_purchase_preference":
            "mse_purchase_preference_status",
    }

    for source_key, fieldname in status_mapping.items():
        raw = extracted.get(
            source_key
        )

        if raw in {
            None,
            "",
        }:
            continue

        value = _business_yes_no_status(
            raw
        )

        if value == "Not Detected":
            continue

        current = getattr(
            tender,
            fieldname,
            None,
        )

        previous = (
            _business_yes_no_status(
                previous_parsed.get(
                    source_key
                )
            )
        )

        if _business_field_is_safe_to_refresh(
            fieldname,
            current,
            previous,
        ):
            setattr(
                tender,
                fieldname,
                value,
            )


def _enrich_business_review_from_documents(
    tender,
):
    """
    Re-evaluate Tender business values using both the primary Tender text
    and downloaded tender-specific supporting documents.

    GeM GTC text is deliberately excluded because it is generic policy text.
    """
    payload = _load_tender_extraction_payload(
        tender
    )

    previous_parsed = payload.get(
        "parsed"
    )

    if not isinstance(
        previous_parsed,
        dict,
    ):
        previous_parsed = {}

    blocks = []

    for page in payload.get(
        "pages"
    ) or []:
        if not isinstance(
            page,
            dict,
        ):
            continue

        value = page.get(
            "text"
        )

        if value:
            blocks.append(
                value
            )

    permitted_classifications = {
        "Buyer ATC",
        "Buyer Specification",
        "Drawing",
        "Specification",
        "MSE / MII Document",
        "Supporting Tender Document",
        "Pre Integrity Pact",
        "Vendor Registration",
        "PQC",
        "Quality",
    }

    for row in (
        tender.tender_source_links
        or []
    ):
        if (
            row.document_read_status
            not in {
                "Read",
                "Read with Warnings",
            }
        ):
            continue

        if (
            row.document_classification
            not in permitted_classifications
        ):
            continue

        if row.document_text:
            blocks.append(
                row.document_text
            )

    if not blocks:
        return {}

    combined_text = "\n\n".join(
        blocks
    )

    extracted = (
        _extract_tender_business_review(
            combined_text,
            previous_parsed,
        )
    )

    _apply_business_review_enrichment(
        tender,
        extracted,
        previous_parsed,
    )

    return extracted


def _build_business_review(tender, parsed):
    """Return a safe, serialisable Tender Review payload."""
    emd_status = _business_yes_no_status(
        parsed.get("emd_required")
    )

    split_status = _business_yes_no_status(
        parsed.get("splitting_applied")
    )

    epbg_percentage = (
        parsed.get("epbg_percentage")
        or getattr(
            tender,
            "epbg_percentage",
            None,
        )
    )

    epbg_duration = (
        parsed.get("epbg_duration_months")
        or getattr(
            tender,
            "epbg_duration_months",
            None,
        )
    )

    return {
        "factory_name":
            parsed.get("factory_name")
            or getattr(
                tender,
                "factory_name",
                None,
            )
            or "Not Detected",

        "nomenclature":
            parsed.get("nomenclature")
            or getattr(
                tender,
                "nomenclature",
                None,
            )
            or "Not Detected",

        "specification_drawing_reference":
            parsed.get(
                "specification_drawing_reference"
            )
            or getattr(
                tender,
                "specification_drawing_reference",
                None,
            )
            or "Not Detected",

        "splitting":
            (
                f"{split_status} - "
                f"{parsed.get('splitting_ratio')}"
                if (
                    split_status == "Yes"
                    and parsed.get("splitting_ratio")
                )
                else split_status
            ),

        "reverse_auction":
            getattr(
                tender,
                "reverse_auction_status",
                None,
            )
            or "Not Detected",

        "tender_type":
            getattr(
                tender,
                "tender_participation_type",
                None,
            )
            or "Not Detected",

        "advance_sample":
            getattr(
                tender,
                "advance_sample_status",
                None,
            )
            or "Not Detected",

        "delivery_period":
            parsed.get(
                "delivery_period_requirement"
            )
            or (
                "Not Detected"
                if (
                    getattr(
                        tender,
                        "delivery_period_requirement",
                        None,
                    )
                    and (
                        "otherwise becomes apparent"
                        in str(
                            getattr(
                                tender,
                                "delivery_period_requirement",
                                "",
                            )
                        ).lower()
                    )
                )
                else (
                    getattr(
                        tender,
                        "delivery_period_requirement",
                        None,
                    )
                    or "Not Detected"
                )
            ),

        "emd_required": emd_status,

        "option_clause":
            getattr(
                tender,
                "option_clause_status",
                None,
            )
            or "Not Detected",

        "documents_required":
            parsed.get(
                "documents_required_from_bidders"
            )
            or getattr(
                tender,
                "documents_required_from_bidders",
                None,
            )
            or "Not Detected",

        "inspection":
            _business_yes_no_status(
                parsed.get(
                    "inspection_required"
                )
            )
            if parsed.get(
                "inspection_required"
            )
            else (
                getattr(
                    tender,
                    "inspection_status",
                    None,
                )
                or "Not Detected"
            ),

        "inspection_type":
            parsed.get("inspection_type")
            or getattr(
                tender,
                "inspection_type",
                None,
            )
            or "Not Detected",

        "inspection_authority":
            parsed.get(
                "inspection_authority"
            )
            or getattr(
                tender,
                "inspection_authority",
                None,
            )
            or "Not Detected",

        "epbg_percentage":
            epbg_percentage
            if epbg_percentage not in {None, ""}
            else "Not Detected",

        "epbg_duration_months":
            epbg_duration
            if epbg_duration not in {None, ""}
            else "Not Detected",

        "mii_purchase_preference":
            getattr(
                tender,
                "mii_purchase_preference_status",
                None,
            )
            or "Not Detected",

        "mse_purchase_preference":
            getattr(
                tender,
                "mse_purchase_preference_status",
                None,
            )
            or "Not Detected",

        "drawing_status":
            getattr(
                tender,
                "drawing_status",
                None,
            )
            or "Not Detected",
    }


def _parse_gem_tender(text):
    """Extract stable GeM Tender header information.

    GeM PDFs contain bilingual labels and formatting varies between
    releases. Separators such as ':' are therefore treated as optional.
    """

    result = {}

    separator = r"\s*(?::|-)?\s*"

    result["bid_number"] = _first_match(
        [
            r"Bid\s*Number"
            + separator
            + r"([A-Z0-9][A-Z0-9/\-]+)",
            r"Bid\s*No\.?"
            + separator
            + r"([A-Z0-9][A-Z0-9/\-]+)",
        ],
        text,
    )

    result["publication_date"] = _first_match(
        [
            r"(?:Dated|Published\s+On)"
            + separator
            + r"(\d{1,2}[-/]\d{1,2}[-/]\d{4})",
        ],
        text,
    )

    result["bid_end"] = _first_match(
        [
            r"Bid\s*End\s*Date\s*/?\s*Time"
            + separator
            + r"("
            r"\d{1,2}[-/]\d{1,2}[-/]\d{4}"
            r"(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?"
            r")",
        ],
        text,
    )

    result["bid_opening"] = _first_match(
        [
            r"Bid\s*Opening\s*Date\s*/?\s*Time"
            + separator
            + r"("
            r"\d{1,2}[-/]\d{1,2}[-/]\d{4}"
            r"(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?"
            r")",
        ],
        text,
    )

    result["ministry"] = _first_match(
        [
            r"Ministry\s*/?\s*State\s*Name"
            + separator
            + r"([^\n]+)",
            r"Ministry"
            + separator
            + r"([^\n]+)",
        ],
        text,
    )

    result["department"] = _first_match(
        [
            r"Department\s*Name"
            + separator
            + r"([^\n]+)",
            r"Department"
            + separator
            + r"([^\n]+)",
        ],
        text,
    )

    result["organisation"] = _first_match(
        [
            r"Organi[sz]ation\s*Name"
            + separator
            + r"([^\n]+)",
            r"Organi[sz]ation"
            + separator
            + r"([^\n]+)",
        ],
        text,
    )

    result["office"] = _first_match(
        [
            r"Office\s*Name"
            + separator
            + r"([^\n]+)",
        ],
        text,
    )

    result["total_quantity"] = _first_match(
        [
            r"Total\s*Quantity"
            + separator
            + r"([\d,.]+)",
        ],
        text,
    )

    result["item_category"] = _first_match(
        [
            r"Item\s*Category"
            + separator
            + r"([^\n]+)",
        ],
        text,
    )

    result["bid_validity"] = _first_match(
        [
            r"Bid\s*Offer\s*Validity"
            r"(?:\s*\(From\s*End\s*Date\))?"
            + separator
            + r"([^\n]+)",
        ],
        text,
    )

    result["emd_amount"] = _first_match(
        [
            r"EMD\s*Amount"
            + separator
            + r"(?:INR|Rs\.?)?\s*([\d,]+(?:\.\d+)?)",
        ],
        text,
    )

    result["emd_required"] = _first_match(
        [
            r"EMD\s*Detail.*?"
            r"Required"
            + separator
            + r"(Yes|No)",
        ],
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )

    # Some GeM formats show EMD Amount without an explicit
    # "Required Yes" row. A positive amount deterministically
    # means EMD is required.
    if (
        not result["emd_required"]
        and result["emd_amount"]
    ):
        try:
            emd_value = float(
                result["emd_amount"].replace(",", "")
            )

            if emd_value > 0:
                result["emd_required"] = "Yes"
        except ValueError:
            pass

    result["epbg_percentage"] = _first_match(
        [
            r"ePBG\s*Percentage\s*\(%\)"
            + separator
            + r"([\d.]+)",
        ],
        text,
    )

    result["epbg_duration_months"] = _first_match(
        [
            r"Duration\s*of\s*ePBG"
            r"\s*required\s*\(Months\)\.?"
            + separator
            + r"([\d.]+)",
        ],
        text,
    )

    result["splitting_applied"] = _first_match(
        [
            r"Splitting\s*Applied"
            + separator
            + r"(Yes|No)",
        ],
        text,
    )

    result["splitting_ratio"] = _first_match(
        [
            r"Split\s*Criteria.*?"
            r"(\d+(?:\s*:\s*\d+)+)",
        ],
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )

    business_review = _extract_tender_business_review(
        text,
        result,
    )

    result.update(
        business_review
    )

    return result

def _parse_date(value):
    if not value:
        return None

    for pattern in (
        "%d-%m-%Y",
        "%d/%m/%Y",
    ):
        try:
            return datetime.strptime(
                value,
                pattern,
            ).date().isoformat()
        except ValueError:
            continue

    return None


def _parse_datetime(value):
    if not value:
        return None

    for pattern in (
        "%d-%m-%Y %H:%M:%S",
        "%d-%m-%Y %H:%M",
        "%d/%m/%Y %H:%M:%S",
        "%d/%m/%Y %H:%M",
    ):
        try:
            return datetime.strptime(
                value,
                pattern,
            ).strftime("%Y-%m-%d %H:%M:%S")
        except ValueError:
            continue

    return None


def _build_summary(parsed, links, page_count):
    lines = [
        f"PDF Pages: {page_count}",
        f"Tender / Bid Reference: {parsed.get('bid_number') or 'Not detected'}",
        f"Publication Date: {parsed.get('publication_date') or 'Not detected'}",
        f"Bid Closing: {parsed.get('bid_end') or 'Not detected'}",
        f"Bid Opening: {parsed.get('bid_opening') or 'Not detected'}",
        f"Bid Validity: {parsed.get('bid_validity') or 'Not detected'}",
        f"Ministry: {parsed.get('ministry') or 'Not detected'}",
        f"Department: {parsed.get('department') or 'Not detected'}",
        f"Organisation: {parsed.get('organisation') or 'Not detected'}",
        f"Office: {parsed.get('office') or 'Not detected'}",
        f"Item Category: {parsed.get('item_category') or 'Not detected'}",
        f"Total Quantity: {parsed.get('total_quantity') or 'Not detected'}",
        f"EMD Required: {parsed.get('emd_required') or 'Not detected'}",
        f"EMD Amount: {parsed.get('emd_amount') or 'Not detected'}",
        f"ePBG Percentage: {parsed.get('epbg_percentage') or 'Not detected'}",
        f"ePBG Duration (Months): {parsed.get('epbg_duration_months') or 'Not detected'}",
        f"Splitting Applied: {parsed.get('splitting_applied') or 'Not detected'}",
        f"Splitting Ratio: {parsed.get('splitting_ratio') or 'Not detected'}",
        f"Hyperlinks Discovered: {len(links)}",
    ]

    return "\n".join(lines)

def _build_warnings(parsed, pages):
    warnings = []

    extracted_text = "\n".join(
        page["text"]
        for page in pages
    ).strip()

    if not extracted_text:
        warnings.append(
            "No machine-readable text was found in the uploaded PDF. "
            "This document may be scanned and may require OCR/manual review."
        )

    if not parsed.get("bid_number"):
        warnings.append(
            "Tender / Bid reference could not be detected automatically."
        )

    if not parsed.get("bid_end"):
        warnings.append(
            "Bid submission deadline could not be detected automatically."
        )

    return warnings


def _is_allowed_gem_hostname(hostname):
    hostname = (hostname or "").lower().rstrip(".")

    if hostname in ALLOWED_GEM_HOSTS:
        return True

    return hostname.endswith(".gem.gov.in")


def _assert_public_ip(hostname):
    try:
        answers = socket.getaddrinfo(
            hostname,
            443,
            proto=socket.IPPROTO_TCP,
        )
    except socket.gaierror as exc:
        frappe.throw(
            _("Unable to resolve Tender link host {0}: {1}").format(
                hostname,
                exc,
            )
        )

    addresses = {
        answer[4][0]
        for answer in answers
    }

    for address in addresses:
        ip = ipaddress.ip_address(address)

        if (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_multicast
            or ip.is_reserved
            or ip.is_unspecified
        ):
            frappe.throw(
                _(
                    "Tender link resolved to a non-public address "
                    "and was blocked."
                )
            )


def _validate_download_url(url):
    parsed = urlparse(url)

    if parsed.scheme.lower() != "https":
        frappe.throw(
            _("Only HTTPS Tender hyperlinks can be downloaded automatically.")
        )

    hostname = (parsed.hostname or "").lower()

    if not _is_allowed_gem_hostname(hostname):
        frappe.throw(
            _(
                "Automatic Tender download is restricted to approved "
                "GeM domains. Host blocked: {0}"
            ).format(hostname or _("Unknown"))
        )

    _assert_public_ip(hostname)

    return hostname


def _filename_from_response(response, final_url):
    """Derive a safe filename for a downloaded Tender document."""

    from urllib.parse import unquote

    content_disposition = (
        response.headers.get("Content-Disposition")
        or ""
    )

    filename = None

    # RFC-style filename*=UTF-8''...
    match = re.search(
        r"filename\*\s*=\s*UTF-8''([^;]+)",
        content_disposition,
        flags=re.IGNORECASE,
    )

    if match:
        filename = unquote(
            match.group(1).strip().strip('"')
        )

    if not filename:
        match = re.search(
            r'filename\s*=\s*"?([^";]+)"?',
            content_disposition,
            flags=re.IGNORECASE,
        )

        if match:
            filename = match.group(1).strip()

    if not filename:
        filename = Path(
            urlparse(final_url).path
        ).name

    content_type = (
        response.headers.get("Content-Type")
        or ""
    ).split(";")[0].strip().lower()

    extension_by_type = {
        "application/pdf": ".pdf",
        "application/msword": ".doc",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            ".docx",
        "application/vnd.ms-excel": ".xls",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            ".xlsx",
    }

    expected_extension = extension_by_type.get(
        content_type
    )

    if not filename:
        filename = "tender_linked_document"

    # Remove path/control characters.
    filename = re.sub(
        r'[\\/:*?"<>|\x00-\x1f]+',
        "_",
        filename,
    ).strip()

    extension = Path(filename).suffix.lower()

    if (
        not extension
        and expected_extension
    ):
        filename += expected_extension
        extension = expected_extension

    if (
        extension
        and extension
        not in ALLOWED_DOWNLOAD_EXTENSIONS
    ):
        raise RuntimeError(
            f"Linked file type {extension} is not approved."
        )

    if (
        content_type
        and content_type != "application/octet-stream"
        and expected_extension is None
    ):
        raise RuntimeError(
            "Linked Tender response type is not approved: "
            f"{content_type}"
        )

    return filename


def _download_allowed_url(url):
    import requests

    current_url = url

    for redirect_index in range(MAX_REDIRECTS + 1):
        _validate_download_url(current_url)

        response = requests.get(
            current_url,
            stream=True,
            timeout=(5, 20),
            allow_redirects=False,
            headers={
                "User-Agent": (
                    "PEPL-Sales-Tender-Ingestion/1.0"
                ),
            },
        )

        if response.status_code in {
            301,
            302,
            303,
            307,
            308,
        }:
            if redirect_index >= MAX_REDIRECTS:
                raise RuntimeError(
                    "Maximum Tender-link redirects exceeded."
                )

            location = response.headers.get("Location")

            if not location:
                raise RuntimeError(
                    "Tender-link redirect had no Location header."
                )

            from urllib.parse import urljoin

            current_url = urljoin(
                current_url,
                location,
            )
            continue

        response.raise_for_status()

        content_length = response.headers.get(
            "Content-Length"
        )

        if (
            content_length
            and int(content_length) > MAX_LINK_DOWNLOAD_BYTES
        ):
            raise RuntimeError(
                "Linked Tender file exceeds the automatic "
                "download size limit."
            )

        content = bytearray()

        for chunk in response.iter_content(
            chunk_size=64 * 1024
        ):
            if not chunk:
                continue

            content.extend(chunk)

            if len(content) > MAX_LINK_DOWNLOAD_BYTES:
                raise RuntimeError(
                    "Linked Tender file exceeds the automatic "
                    "download size limit."
                )

        filename = _filename_from_response(
            response,
            current_url,
        )

        return {
            "content": bytes(content),
            "filename": filename,
            "final_url": current_url,
            "content_type": response.headers.get(
                "Content-Type",
                "",
            ),
        }

    raise RuntimeError(
        "Unable to retrieve linked Tender document."
    )


def _attach_content(
    *,
    content,
    filename,
    tender_name,
    attached_to_field=None,
):
    values = {
        "doctype": "File",
        "file_name": filename,
        "attached_to_doctype": "PEPL Tender",
        "attached_to_name": tender_name,
        "content": content,
        "is_private": 1,
    }

    if attached_to_field:
        values["attached_to_field"] = attached_to_field

    file_doc = frappe.get_doc(values)
    file_doc.insert(ignore_permissions=True)

    return file_doc.file_url


def _populate_source_links(tender, links):
    """
    Synchronize discovered Tender URLs without destroying previously
    downloaded/read supporting-document state.

    Existing rows are matched by Source URL and preserved in place.
    Newly discovered URLs are appended as new rows. URLs no longer
    present in a subsequent PDF read are retained as historical Tender
    evidence rather than silently deleting attachments.
    """

    existing_by_url = {}

    for row in (
        tender.tender_source_links
        or []
    ):
        source_url = (
            row.source_url
            or ""
        ).strip()

        if (
            source_url
            and source_url not in existing_by_url
        ):
            existing_by_url[
                source_url
            ] = row

    seen_urls = set()

    for link in links:
        source_url = (
            link.get("url")
            or ""
        ).strip()

        if (
            not source_url
            or source_url in seen_urls
        ):
            continue

        seen_urls.add(
            source_url
        )

        status = (
            "Discovered"
            if (
                link.get("link_type")
                == "GeM Document"
            )
            else "Manual Review Required"
        )

        existing = existing_by_url.get(
            source_url
        )

        if existing:
            # Refresh discovery metadata only. Never reset retrieval,
            # download, classification, read status, extracted text,
            # page count or existing processing remarks.
            if link.get("host"):
                existing.source_host = (
                    link["host"]
                )

            if link.get("link_type"):
                existing.link_type = (
                    link["link_type"]
                )

            if not existing.retrieval_status:
                existing.retrieval_status = (
                    status
                )

            if not existing.remarks:
                existing.remarks = (
                    f"Discovered on PDF page "
                    f"{link.get('page') or '-'}."
                )

            continue

        tender.append(
            "tender_source_links",
            {
                "source_url": source_url,
                "source_host": (
                    link.get("host")
                    or ""
                ),
                "link_type": (
                    link.get("link_type")
                    or ""
                ),
                "retrieval_status": status,
                "remarks": (
                    f"Discovered on PDF page "
                    f"{link.get('page') or '-'}."
                ),
            },
        )



def _set_safe_extracted_fields(tender, parsed):
    """
    Populate deterministic Tender identification and the dedicated
    Tender Business Review fields.

    Customer, Sector, Items and operational commercial fields are not
    inferred or overwritten here. Existing user-reviewed business
    values are also preserved.
    """

    if (
        not tender.nit_number
        and parsed.get("bid_number")
    ):
        tender.nit_number = parsed["bid_number"]

    if (
        not tender.publication_date
        and parsed.get("publication_date")
    ):
        tender.publication_date = _parse_date(
            parsed["publication_date"]
        )

    if (
        not tender.bid_submission_deadline
        and parsed.get("bid_end")
    ):
        tender.bid_submission_deadline = _parse_datetime(
            parsed["bid_end"]
        )

    if (
        not tender.bid_opening_date
        and parsed.get("bid_opening")
    ):
        tender.bid_opening_date = _parse_datetime(
            parsed["bid_opening"]
        )

    def set_if_unreviewed(
        fieldname,
        value,
        *,
        empty_values=None,
    ):
        if value in {None, ""}:
            return

        if empty_values is None:
            empty_values = {
                None,
                "",
                "Not Detected",
            }

        current = getattr(
            tender,
            fieldname,
            None,
        )

        if current in empty_values:
            setattr(
                tender,
                fieldname,
                value,
            )

    set_if_unreviewed(
        "factory_name",
        parsed.get("factory_name"),
    )

    # Primary Tender extraction owns Nomenclature. The value "of"
    # is a known stale automatic fragment produced by an older
    # supporting-document extraction rule and is safe to repair.
    set_if_unreviewed(
        "nomenclature",
        parsed.get("nomenclature"),
        empty_values={
            None,
            "",
            "Not Detected",
            "of",
        },
    )

    set_if_unreviewed(
        "specification_drawing_reference",
        parsed.get(
            "specification_drawing_reference"
        ),
    )

    set_if_unreviewed(
        "tender_participation_type",
        parsed.get(
            "tender_participation_type"
        ),
    )

    set_if_unreviewed(
        "reverse_auction_status",
        _business_yes_no_status(
            parsed.get("reverse_auction")
        ),
    )

    set_if_unreviewed(
        "advance_sample_status",
        _business_yes_no_status(
            parsed.get("advance_sample")
        ),
    )

    set_if_unreviewed(
        "delivery_period_requirement",
        parsed.get(
            "delivery_period_requirement"
        ),
    )

    set_if_unreviewed(
        "option_clause_status",
        _business_yes_no_status(
            parsed.get("option_clause")
        ),
    )

    set_if_unreviewed(
        "documents_required_from_bidders",
        parsed.get(
            "documents_required_from_bidders"
        ),
    )

    emd_status = _business_yes_no_status(
        parsed.get("emd_required")
    )

    emd_detection = {
        "Yes": "Required",
        "No": "Not Required",
    }.get(
        emd_status,
        "Not Detected",
    )

    set_if_unreviewed(
        "emd_detection_status",
        emd_detection,
    )

    set_if_unreviewed(
        "inspection_status",
        _business_yes_no_status(
            parsed.get("inspection_required")
        ),
    )

    set_if_unreviewed(
        "inspection_type",
        parsed.get("inspection_type"),
    )

    set_if_unreviewed(
        "inspection_authority",
        parsed.get(
            "inspection_authority"
        ),
    )

    epbg_percentage = parsed.get(
        "epbg_percentage"
    )

    if epbg_percentage not in {
        None,
        "",
    }:
        try:
            set_if_unreviewed(
                "epbg_percentage",
                float(epbg_percentage),
                empty_values={
                    None,
                    "",
                    0,
                    0.0,
                },
            )
        except (TypeError, ValueError):
            pass

    epbg_duration = parsed.get(
        "epbg_duration_months"
    )

    if epbg_duration not in {
        None,
        "",
    }:
        try:
            set_if_unreviewed(
                "epbg_duration_months",
                int(
                    float(
                        epbg_duration
                    )
                ),
                empty_values={
                    None,
                    "",
                    0,
                },
            )
        except (TypeError, ValueError):
            pass

    set_if_unreviewed(
        "mii_purchase_preference_status",
        _business_yes_no_status(
            parsed.get(
                "mii_purchase_preference"
            )
        ),
    )

    set_if_unreviewed(
        "mse_purchase_preference_status",
        _business_yes_no_status(
            parsed.get(
                "mse_purchase_preference"
            )
        ),
    )

    # A drawing/specification reference does not prove that a
    # downloadable drawing exists. Keep this separate until supporting
    # Tender documents are actually processed.
    if (
        parsed.get(
            "specification_drawing_reference"
        )
        and getattr(
            tender,
            "drawing_status",
            None,
        )
        in {
            None,
            "",
            "Not Detected",
        }
    ):
        tender.drawing_status = (
            "Manual Review Required"
        )


@frappe.whitelist()
def read_tender_pdf(tender_name):
    tender = frappe.get_doc(
        "PEPL Tender",
        tender_name,
    )

    tender.check_permission("write")

    if tender.docstatus != 0:
        frappe.throw(
            _("Tender PDF can be read only while the Tender is in Draft.")
        )

    if not tender.tender_source_pdf:
        frappe.throw(
            _("Upload Original Tender PDF first.")
        )

    if not tender.tender_source_pdf.lower().endswith(
        ".pdf"
    ):
        frappe.throw(
            _("Original Tender file must be a PDF.")
        )

    try:
        content = _get_attached_file_content(
            tender.tender_source_pdf
        )

        if len(content) > MAX_SOURCE_PDF_BYTES:
            frappe.throw(
                _("Tender PDF exceeds the 25 MB reading limit.")
            )

        reader = PdfReader(
            BytesIO(content)
        )

        pages = _extract_pdf_text(reader)
        links = _extract_pdf_links(reader)

        full_text = "\n\n".join(
            page["text"]
            for page in pages
        )

        parsed = _parse_gem_tender(
            full_text
        )

        warnings = _build_warnings(
            parsed,
            pages,
        )

        extraction = {
            "engine_version": "1.1",
            "source_file": tender.tender_source_pdf,
            "page_count": len(reader.pages),
            "parsed": parsed,
            "links": links,
            "pages": pages,
            "warnings": warnings,
        }

        _populate_source_links(
            tender,
            links,
        )

        _set_safe_extracted_fields(
            tender,
            parsed,
        )

        # Refresh business-review fields from the current extraction.
        # This also replaces recognised false-positive values generated
        # by older parser versions while preserving sensible user values.
        _apply_business_review_enrichment(
            tender,
            parsed,
            {},
        )

        tender.tender_ingestion_summary = (
            _build_summary(
                parsed,
                links,
                len(reader.pages),
            )
        )

        tender.tender_extraction_json = (
            json.dumps(
                extraction,
                ensure_ascii=False,
                indent=2,
            )
        )

        tender.tender_ingestion_warnings = (
            "\n".join(warnings)
        )

        tender.tender_ingestion_status = (
            "Read with Warnings"
            if warnings
            else "Read - Review Required"
        )

        tender.tender_ingestion_last_run = (
            now_datetime()
        )

        tender.save()

        return {
            "status": tender.tender_ingestion_status,
            "page_count": len(reader.pages),
            "link_count": len(links),
            "parsed": parsed,
            "business_review": _build_business_review(
                tender,
                parsed,
            ),
            "warning_count": len(warnings),
            "warnings": warnings,
        }

    except Exception as exc:
        frappe.log_error(
            frappe.get_traceback(),
            "PEPL Tender Automatic Reading",
        )

        tender.db_set(
            "tender_ingestion_status",
            "Failed",
            update_modified=False,
        )

        tender.db_set(
            "tender_ingestion_warnings",
            str(exc),
            update_modified=True,
        )

        raise



@frappe.whitelist()
def download_discovered_gem_documents(
    tender_name,
):
    tender = frappe.get_doc(
        "PEPL Tender",
        tender_name,
    )

    tender.check_permission("write")

    if tender.docstatus != 0:
        frappe.throw(
            _("Linked Tender files can be retrieved only in Draft.")
        )

    downloaded = 0
    portal_reference = 0
    manual_download = 0
    unavailable = 0
    failed = 0
    skipped = 0

    for row in tender.tender_source_links or []:
        if row.link_type != "GeM Document":
            skipped += 1
            continue

        if row.downloaded_file:
            skipped += 1
            continue

        if row.retrieval_status in {
            "Portal / Reference Link",
            "Manual Download Required",
            "Unavailable",
        }:
            skipped += 1
            continue

        try:
            result = _download_allowed_url(
                row.source_url
            )

            file_url = _attach_content(
                content=result["content"],
                filename=result["filename"],
                tender_name=tender.name,
            )

            row.downloaded_file = file_url
            row.retrieval_status = "Downloaded"
            row.document_classification = (
                _classify_supporting_document(
                    result["filename"],
                    "",
                    source_url=row.source_url,
                    preferred_classification=(
                        row.document_classification
                        or None
                    ),
                )
            )
            row.remarks = (
                "Downloaded securely from approved GeM host."
            )

            if (
                row.document_classification == "Drawing"
            ):
                tender.drawing_status = "Available"
                tender.drawing_supporting_file = (
                    file_url
                )

            downloaded += 1

        except Exception as exc:
            message = str(exc)
            lower_message = message.lower()

            if (
                "text/html" in lower_message
                or "browser/session" in lower_message
                or "login" in lower_message
                or "authentication" in lower_message
                or "authorisation" in lower_message
                or "authorization" in lower_message
            ):
                row.retrieval_status = (
                    "Manual Download Required"
                )
                row.remarks = (
                    "GeM requires browser/session interaction "
                    "for this resource. Open the URL manually "
                    "and attach the supporting document if it "
                    "is required for this Tender."
                )
                manual_download += 1

            elif (
                "not approved" in lower_message
                or "unsupported content" in lower_message
            ):
                row.retrieval_status = (
                    "Portal / Reference Link"
                )
                row.remarks = (
                    "The GeM URL resolved to a reference or "
                    "non-downloadable portal resource rather "
                    "than an approved Tender attachment."
                )
                portal_reference += 1

            elif (
                "404" in lower_message
                or "410" in lower_message
                or "not found" in lower_message
                or "no longer available" in lower_message
            ):
                row.retrieval_status = "Unavailable"
                row.remarks = (
                    "The GeM resource is currently unavailable. "
                    + message[:350]
                )
                unavailable += 1

            else:
                row.retrieval_status = "Failed"
                row.remarks = message[:500]
                failed += 1

    tender.save()

    return {
        "downloaded": downloaded,
        "portal_reference": portal_reference,
        "manual_download": manual_download,
        "unavailable": unavailable,
        "failed": failed,
        "skipped": skipped,
    }


def _extract_primary_document_titles(tender):
    """Return GeM Additional Scope document titles in source order."""

    if not tender.tender_extraction_json:
        return []

    try:
        extraction = json.loads(
            tender.tender_extraction_json
        )
    except Exception:
        return []

    pages = extraction.get(
        "pages",
        [],
    )

    canonical_titles = [
        "Drawing",
        "Specification",
        "Quality",
        "PQC",
        "Vendor Registration",
        "Pre Integrity Pact",
    ]

    keyword_map = {
        "Drawing": [
            "DRAWING",
        ],
        "Specification": [
            "SPECIFICATION",
        ],
        "Quality": [
            "QUALITY",
        ],
        "PQC": [
            "PQC",
        ],
        "Vendor Registration": [
            "VRAF",
        ],
        "Pre Integrity Pact": [
            "PRE INTEGRITY PACT",
        ],
    }

    for page in pages:
        page_text = (
            page.get("text")
            or ""
        )

        upper_text = page_text.upper()

        if "DOCUMENT TITLE" not in upper_text:
            continue

        found = []

        for title in canonical_titles:
            positions = []

            for keyword in keyword_map[title]:
                position = upper_text.find(
                    keyword.upper()
                )

                if position >= 0:
                    positions.append(
                        position
                    )

            if positions:
                found.append(
                    (
                        min(positions),
                        title,
                    )
                )

        found.sort(
            key=lambda value: value[0]
        )

        return [
            title
            for _, title in found
        ]

    return []



def _classify_supporting_document(
    filename,
    text_content,
    source_url=None,
    preferred_classification=None,
):
    """Classify a Tender document using strong source metadata first."""
    if preferred_classification:
        return preferred_classification

    source_url = (
        source_url
        or ""
    ).lower()

    filename_lower = (
        filename
        or ""
    ).lower()

    text_lower = (
        text_content
        or ""
    ).lower()

    if (
        "admin.gem.gov.in" in source_url
        or "new-gtc" in filename_lower
        or "gem-gtc" in filename_lower
    ):
        return "GeM GTC"

    if (
        "/specificationdocument/" in source_url
        or "specificationdocument" in source_url
    ):
        return "Buyer Specification"

    # GeM fulfilment ATC links are commonly PDF files, not only DOCX.
    # Detect them before generic content classification.
    if (
        "fulfilment.gem.gov.in" in source_url
        and (
            "atc" in filename_lower
            or "atc" in source_url
            or (
                "compliance statement for the tender"
                in text_lower
            )
        )
    ):
        return "Buyer ATC"

    if (
        "atc - compliance statement"
        in text_lower
        or (
            "gem bid no."
            in text_lower
            and "clause description"
            in text_lower
            and "compliance by bidder"
            in text_lower
        )
    ):
        return "Buyer ATC"

    value = (
        f"{filename or ''}\n"
        f"{text_content or ''}"
    ).lower()

    # Source-specific and document-purpose classifications come before
    # generic keywords that may occur inside a large ATC bundle.
    rules = [
        (
            "Buyer ATC",
            [
                "buyer uploaded atc",
                "buyer added bid specific atc",
                "additional terms and conditions",
                "atc compliance statement",
            ],
        ),
        (
            "Pre Integrity Pact",
            [
                "pre integrity pact",
                "integrity pact",
                "annexure 5b",
            ],
        ),
        (
            "Vendor Registration",
            [
                "vraf",
                "vendor application",
                "vendor registration",
            ],
        ),
        (
            "PQC",
            [
                "pqc",
                "capacity verification",
                "annexure b",
            ],
        ),
        (
            "Quality",
            [
                "annexure q",
                "quality assurance",
                "quality requirement",
            ],
        ),
        (
            "Buyer Specification",
            [
                "buyer specification",
                "technical specification",
                "specification document",
            ],
        ),
        (
            "Drawing",
            [
                "drawing",
                "drg.",
                "drg ",
            ],
        ),
        (
            "Specification",
            [
                "specification",
                "annexure a",
            ],
        ),
        (
            "MSE / MII Document",
            [
                "mse",
                "make in india",
                "mii",
            ],
        ),
    ]

    for classification, keywords in rules:
        for keyword in keywords:
            if keyword in value:
                return classification

    return "Supporting Tender Document"


def _extract_pdf_supporting_document(content):
    reader = PdfReader(
        BytesIO(content)
    )

    pages = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        try:
            page_text = (
                page.extract_text()
                or ""
            )
        except Exception:
            page_text = ""

        page_text = _normalise_multiline_text(
            page_text
        )

        if page_text:
            pages.append(
                f"[Page {page_number}]\n"
                f"{page_text}"
            )

    return {
        "page_count": len(reader.pages),
        "text": "\n\n".join(pages),
    }


def _extract_docx_supporting_document(content):
    from docx import Document as WordDocument

    document = WordDocument(
        BytesIO(content)
    )

    blocks = []

    for paragraph in document.paragraphs:
        value = _clean_text(
            paragraph.text
        )

        if value:
            blocks.append(value)

    for table_number, table in enumerate(
        document.tables,
        start=1,
    ):
        blocks.append(
            f"[Table {table_number}]"
        )

        for row in table.rows:
            values = []

            for cell in row.cells:
                value = _clean_text(
                    cell.text
                )

                values.append(value)

            if any(values):
                blocks.append(
                    " | ".join(values)
                )

    normal_text = "\n".join(
        blocks
    ).strip()

    if normal_text:
        return {
            "page_count": 0,
            "text": normal_text,
        }

    # Fallback for Word textboxes/shapes that python-docx
    # does not expose through document.paragraphs.
    import zipfile
    import xml.etree.ElementTree as ET

    xml_blocks = []

    word_namespace = (
        "{http://schemas.openxmlformats.org/"
        "wordprocessingml/2006/main}"
    )

    with zipfile.ZipFile(
        BytesIO(content)
    ) as archive:
        xml_names = [
            name
            for name in archive.namelist()
            if (
                name == "word/document.xml"
                or name.startswith(
                    "word/header"
                )
                or name.startswith(
                    "word/footer"
                )
                or name in {
                    "word/footnotes.xml",
                    "word/endnotes.xml",
                }
            )
        ]

        for xml_name in xml_names:
            try:
                root = ET.fromstring(
                    archive.read(
                        xml_name
                    )
                )
            except Exception:
                continue

            values = []

            for node in root.iter(
                word_namespace + "t"
            ):
                value = _clean_text(
                    node.text
                )

                if value:
                    values.append(value)

            if values:
                xml_blocks.append(
                    "\n".join(values)
                )

    return {
        "page_count": 0,
        "text": "\n\n".join(
            xml_blocks
        ),
    }



def _read_supporting_file(file_url):
    file_name = frappe.db.get_value(
        "File",
        {
            "file_url": file_url,
        },
        "name",
    )

    if not file_name:
        raise RuntimeError(
            f"Supporting File not found: {file_url}"
        )

    file_doc = frappe.get_doc(
        "File",
        file_name,
    )

    content = file_doc.get_content()

    filename = (
        file_doc.file_name
        or Path(file_url).name
    )

    extension = Path(
        filename
    ).suffix.lower()

    if extension == ".pdf":
        result = (
            _extract_pdf_supporting_document(
                content
            )
        )

    elif extension == ".docx":
        result = (
            _extract_docx_supporting_document(
                content
            )
        )

    elif extension in {
        ".doc",
        ".xls",
        ".xlsx",
    }:
        return {
            "filename": filename,
            "extension": extension,
            "page_count": 0,
            "text": "",
            "status":
                "Manual Review Required",
            "warning": (
                f"{extension} automatic text "
                "extraction is not enabled."
            ),
        }

    else:
        return {
            "filename": filename,
            "extension": extension,
            "page_count": 0,
            "text": "",
            "status":
                "Manual Review Required",
            "warning": (
                "Unsupported supporting "
                f"document type: {extension or 'unknown'}"
            ),
        }

    extracted_text = (
        result.get("text")
        or ""
    ).strip()

    page_count = (
        result.get("page_count")
        or 0
    )

    if not extracted_text:
        return {
            "filename": filename,
            "extension": extension,
            "page_count": page_count,
            "text": "",
            "status":
                "Read with Warnings",
            "warning": (
                "No machine-readable text was "
                "found in this supporting document."
            ),
        }

    if extension == ".pdf":
        minimum_expected_text = max(
            100,
            page_count * 80,
        )

        if len(extracted_text) < minimum_expected_text:
            return {
                "filename": filename,
                "extension": extension,
                "page_count": page_count,
                "text": extracted_text,
                "status":
                    "Read with Warnings",
                "warning": (
                    "Very little machine-readable text was found. "
                    "This supporting PDF may contain scanned pages "
                    "or drawings and requires visual/manual review."
                ),
            }

    if (
        extension == ".docx"
        and len(extracted_text) < 50
    ):
        return {
            "filename": filename,
            "extension": extension,
            "page_count": page_count,
            "text": extracted_text,
            "status":
                "Read with Warnings",
            "warning": (
                "Very little machine-readable text was found "
                "inside this Word document. The file may contain "
                "images, drawing objects or unsupported embedded "
                "content and requires visual/manual review."
            ),
        }

    return {
        "filename": filename,
        "extension": extension,
        "page_count": page_count,
        "text": extracted_text,
        "status": "Read",
        "warning": "",
    }


@frappe.whitelist()

def read_supporting_tender_documents(
    tender_name,
):
    tender = frappe.get_doc(
        "PEPL Tender",
        tender_name,
    )

    tender.check_permission("write")

    if tender.docstatus != 0:
        frappe.throw(
            _(
                "Supporting Tender documents can "
                "be read only while the Tender is Draft."
            )
        )

    read_count = 0
    warning_count = 0
    manual_count = 0
    failed_count = 0
    skipped_count = 0

    primary_document_titles = (
        _extract_primary_document_titles(
            tender
        )
    )

    direct_document_rows = [
        row
        for row in (
            tender.tender_source_links
            or []
        )
        if (
            "/bidding/bid/documentdownload/"
            in (row.source_url or "")
        )
    ]

    preferred_classification = {}

    for index, row in enumerate(
        direct_document_rows
    ):
        if index < len(
            primary_document_titles
        ):
            preferred_classification[
                row.name
            ] = primary_document_titles[
                index
            ]

    for row in tender.tender_source_links or []:
        if not row.downloaded_file:
            skipped_count += 1
            continue

        try:
            result = _read_supporting_file(
                row.downloaded_file
            )

            row.document_page_count = (
                result.get("page_count")
                or 0
            )

            row.document_text = (
                result.get("text")
                or ""
            )

            row.document_read_status = (
                result.get("status")
                or "Failed"
            )

            row.document_classification = (
                _classify_supporting_document(
                    result.get("filename"),
                    result.get("text"),
                    source_url=row.source_url,
                    preferred_classification=(
                        preferred_classification.get(
                            row.name
                        )
                    ),
                )
            )

            warning = (
                result.get("warning")
                or ""
            )

            if (
                warning
                and warning not in (
                    row.remarks
                    or ""
                )
            ):
                row.remarks = (
                    (
                        row.remarks
                        + "\n"
                    )
                    if row.remarks
                    else ""
                ) + warning

            if (
                row.document_read_status
                == "Read"
            ):
                read_count += 1

            elif (
                row.document_read_status
                == "Read with Warnings"
            ):
                warning_count += 1

            elif (
                row.document_read_status
                == "Manual Review Required"
            ):
                manual_count += 1

            else:
                failed_count += 1

        except Exception as exc:
            row.document_read_status = "Failed"
            row.document_text = ""
            row.document_page_count = 0

            row.remarks = (
                (
                    row.remarks
                    + "\n"
                )
                if row.remarks
                else ""
            ) + str(exc)[:500]

            failed_count += 1

    enriched = (
        _enrich_business_review_from_documents(
            tender
        )
    )

    tender.save()

    return {
        "read": read_count,
        "warnings": warning_count,
        "manual_review": manual_count,
        "failed": failed_count,
        "skipped": skipped_count,
        "business_review_enriched": bool(
            enriched
        ),
    }


@frappe.whitelist()
def mark_tender_extraction_reviewed(
    tender_name,
):
    tender = frappe.get_doc(
        "PEPL Tender",
        tender_name,
    )

    tender.check_permission("write")

    if tender.docstatus != 0:
        frappe.throw(
            _("Extraction review can be completed only in Draft.")
        )

    if tender.tender_ingestion_status not in {
        "Read - Review Required",
        "Read with Warnings",
        "Reviewed",
    }:
        frappe.throw(
            _("Read the Tender PDF before completing extraction review.")
        )

    tender.tender_ingestion_status = "Reviewed"
    tender.save()

    tender.add_comment(
        "Comment",
        text=(
            "Tender automatic extraction reviewed by "
            f"{frappe.session.user}."
        ),
    )

    return {
        "status": "Reviewed",
    }


def _set_docx_defaults(document):
    """Preserve the official PEPL Word-template layout.

    The supplied PEPL letterhead uses a non-standard Word style set.
    Tender generation therefore uses direct paragraph/run formatting
    instead of depending on Normal/Heading/List named styles.
    """

    # Intentionally do not modify:
    # - section margins
    # - header/footer distances
    # - corporate header/footer
    # - named styles
    #
    # The supplied PEPL DOCX remains the layout authority.
    return document


def _add_docx_heading(
    document,
    text,
    level=1,
    *,
    centered=False,
):
    """Add a heading without relying on Word named styles."""

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    sizes = {
        1: 16,
        2: 13,
        3: 11,
    }

    paragraph = document.add_paragraph()

    if centered:
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    paragraph.paragraph_format.space_before = Pt(
        8 if level > 1 else 10
    )
    paragraph.paragraph_format.space_after = Pt(
        5
    )

    run = paragraph.add_run(
        str(text or "")
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(
        sizes.get(level, 11)
    )

    return paragraph


def _add_docx_bullet(
    document,
    text,
):
    """Add a bullet without relying on the List Bullet style."""

    from docx.shared import Pt

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = Pt(
        14
    )
    paragraph.paragraph_format.first_line_indent = Pt(
        -10
    )
    paragraph.paragraph_format.space_after = Pt(
        3
    )

    bullet = paragraph.add_run(
        "\u2022 "
    )
    bullet.font.name = "Arial"
    bullet.font.size = Pt(10)

    content = paragraph.add_run(
        str(text or "")
    )
    content.font.name = "Arial"
    content.font.size = Pt(10)

    return paragraph



def _add_key_value_table(
    document,
    rows,
):
    table = document.add_table(
        rows=0,
        cols=2,
    )

    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = (
            str(value)
            if value not in (None, "")
            else "-"
        )

        if cells[0].paragraphs:
            run = cells[0].paragraphs[0].runs[0]
            run.bold = True

    return table


def _load_word_letterhead():
    from docx import Document as WordDocument

    letterhead_path = Path(
        frappe.get_app_path(
            "pepl_sales",
            "public",
            "letterhead",
            "PEPL_Letterhead_Plain.docx",
        )
    )

    if letterhead_path.exists():
        return WordDocument(
            str(letterhead_path)
        )

    return WordDocument()



def _supporting_document_excerpt(row):
    """Return only useful machine-readable content for the working DOCX.

    The authoritative supporting file always remains attached to ERPNext.
    Warning/scanned/drawing-heavy documents are referenced in the register
    and manual-review section rather than dumping unreliable extracted text.
    """

    text_content = (
        row.document_text
        or ""
    ).strip()

    if not text_content:
        return ""

    classification = (
        row.document_classification
        or "Supporting Tender Document"
    )

    # Never place unreliable extraction into the client-facing
    # editable Tender document.
    if (
        row.document_read_status
        != "Read"
    ):
        return ""

    # Drawings usually extract as layout noise, while the full GeM GTC
    # is intentionally retained only as an authoritative source file.
    if classification in {
        "Drawing",
        "GeM GTC",
    }:
        return ""

    limits = {
        "Buyer Specification": 3000,
        "Specification": 3000,
        "Quality": 1500,
        "PQC": 1500,
        "Vendor Registration": 1800,
        "Pre Integrity Pact": 2200,
        "Buyer ATC": 2500,
        "Supporting Tender Document": 1800,
    }

    limit = limits.get(
        classification,
        1800,
    )

    excerpt = text_content[:limit]

    if len(text_content) > limit:
        excerpt += (
            "\n\n[Excerpt truncated. "
            "Refer to the original supporting file in ERPNext.]"
        )

    return excerpt



def _add_supporting_document_register(
    document,
    tender,
):
    _add_docx_heading(
        document,
        "Supporting Tender Document Register",
        level=2,
    )

    if not tender.tender_source_links:
        document.add_paragraph(
            "No supporting Tender hyperlinks or attachments were recorded."
        )
        return

    table = document.add_table(
        rows=1,
        cols=7,
    )

    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    headings = [
        "No.",
        "Classification",
        "Retrieval",
        "Read Status",
        "Pages",
        "Supporting File",
        "Review Required",
    ]

    for index, heading in enumerate(
        headings
    ):
        table.rows[0].cells[
            index
        ].text = heading

    for row in tender.tender_source_links:
        cells = table.add_row().cells

        review_required = (
            "Yes"
            if (
                row.retrieval_status
                == "Manual Review Required"
                or row.document_read_status
                in {
                    "Read with Warnings",
                    "Manual Review Required",
                    "Failed",
                }
                or not row.downloaded_file
            )
            else "No"
        )

        values = [
            row.idx,
            (
                row.document_classification
                or row.link_type
                or "-"
            ),
            row.retrieval_status or "-",
            row.document_read_status or "Not Read",
            row.document_page_count or 0,
            row.downloaded_file or "-",
            review_required,
        ]

        for index, value in enumerate(
            values
        ):
            cells[index].text = str(
                value
            )


def _add_supporting_document_excerpts(
    document,
    tender,
):
    _add_docx_heading(
        document,
        "Extracted Supporting Document Highlights",
        level=2,
    )

    included = 0

    for row in tender.tender_source_links or []:
        excerpt = (
            _supporting_document_excerpt(
                row
            )
        )

        if not row.downloaded_file:
            continue

        included += 1

        classification = (
            row.document_classification
            or f"Supporting Document {row.idx}"
        )

        _add_docx_heading(
            document,
            classification,
            level=3,
        )

        metadata = document.add_paragraph()

        metadata.add_run(
            "File: "
        ).bold = True

        metadata.add_run(
            row.downloaded_file
            or "-"
        )

        metadata.add_run(
            " | Read Status: "
        ).bold = True

        metadata.add_run(
            row.document_read_status
            or "Not Read"
        )

        if row.document_page_count:
            metadata.add_run(
                " | Pages: "
            ).bold = True

            metadata.add_run(
                str(
                    row.document_page_count
                )
            )

        if excerpt:
            document.add_paragraph(
                excerpt
            )
        else:
            if (
                row.document_read_status
                != "Read"
            ):
                note = (
                    "Machine-readable content is incomplete or "
                    "unreliable. Refer to the original supporting "
                    "file and complete visual/manual review."
                )

            elif classification == "Drawing":
                note = (
                    "Engineering drawing retained as the authoritative "
                    "supporting file. Raw PDF text extraction is not "
                    "included in this working document."
                )

            elif classification == "GeM GTC":
                note = (
                    "Full GeM General Terms and Conditions are retained "
                    "as the authoritative supporting file in ERPNext "
                    "and are not reproduced in this working document."
                )

            else:
                note = (
                    "Refer to the original supporting file in ERPNext."
                )

            document.add_paragraph(
                note
            )

    if not included:
        document.add_paragraph(
            "No machine-readable supporting-document excerpts "
            "were available. Refer to the original Tender attachments."
        )


def _add_manual_review_register(
    document,
    tender,
):
    review_rows = []

    for row in tender.tender_source_links or []:
        requires_review = (
            row.retrieval_status
            == "Manual Review Required"
            or row.document_read_status
            in {
                "Read with Warnings",
                "Manual Review Required",
                "Failed",
            }
            or not row.downloaded_file
        )

        if requires_review:
            review_rows.append(
                row
            )

    _add_docx_heading(
        document,
        "Documents Requiring Visual / Manual Review",
        level=2,
    )

    if not review_rows:
        document.add_paragraph(
            "No supporting documents are currently flagged "
            "for visual/manual review."
        )
        return

    for row in review_rows:
        classification = (
            row.document_classification
            or row.link_type
            or f"Document {row.idx}"
        )

        paragraph = _add_docx_bullet(
            document,
            classification,
        )

        if paragraph.runs:
            paragraph.runs[-1].bold = True

        details = []

        if row.retrieval_status:
            details.append(
                f"Retrieval: {row.retrieval_status}"
            )

        if row.document_read_status:
            details.append(
                f"Read: {row.document_read_status}"
            )

        if row.remarks:
            details.append(
                f"Remarks: {row.remarks}"
            )

        if row.source_url:
            details.append(
                f"Source: {row.source_url}"
            )

        paragraph.add_run(
            " — "
            + " | ".join(details)
        )



def _reviewed_yes_no(value):
    value = str(
        value or ""
    ).strip().lower()

    if value in {
        "yes",
        "y",
        "true",
        "1",
    }:
        return 1

    if value in {
        "no",
        "n",
        "false",
        "0",
    }:
        return 0

    return None


def _get_reviewed_tender_parsed_data(
    tender,
):
    if (
        tender.tender_ingestion_status
        != "Reviewed"
    ):
        frappe.throw(
            _(
                "Tender extraction must be Reviewed "
                "before applying extracted values."
            )
        )

    if not tender.tender_extraction_json:
        frappe.throw(
            _(
                "Tender Extraction Data is not available."
            )
        )

    try:
        extraction = json.loads(
            tender.tender_extraction_json
        )
    except Exception:
        frappe.throw(
            _(
                "Tender Extraction Data contains invalid JSON."
            )
        )

    parsed = extraction.get(
        "parsed"
    ) or {}

    if not parsed:
        frappe.throw(
            _(
                "No parsed Tender information is available."
            )
        )

    return parsed


@frappe.whitelist()
def get_reviewed_tender_extraction_preview(
    tender_name,
):
    tender = frappe.get_doc(
        "PEPL Tender",
        tender_name,
    )

    tender.check_permission("read")

    parsed = (
        _get_reviewed_tender_parsed_data(
            tender
        )
    )

    extracted_emd_required = (
        _reviewed_yes_no(
            parsed.get(
                "emd_required"
            )
        )
    )

    extracted_splitting = (
        _reviewed_yes_no(
            parsed.get(
                "splitting_applied"
            )
        )
    )

    return {
        "current": {
            "publication_date":
                tender.publication_date,
            "bid_submission_deadline":
                tender.bid_submission_deadline,
            "bid_opening_date":
                tender.bid_opening_date,
            "emd_required":
                tender.emd_required,
            "emd_amount":
                tender.emd_amount,
            "bid_securing_declaration":
                tender.bid_securing_declaration,
            "emd_mode":
                tender.emd_mode,
            "splitting_applicable":
                tender.splitting_applicable,
        },
        "extracted": {
            "publication_date":
                _parse_date(
                    parsed.get(
                        "publication_date"
                    )
                )
                if parsed.get(
                    "publication_date"
                )
                else None,
            "bid_submission_deadline":
                _parse_datetime(
                    parsed.get(
                        "bid_end"
                    )
                )
                if parsed.get(
                    "bid_end"
                )
                else None,
            "bid_opening_date":
                _parse_datetime(
                    parsed.get(
                        "bid_opening"
                    )
                )
                if parsed.get(
                    "bid_opening"
                )
                else None,
            "emd_required":
                extracted_emd_required,
            "emd_amount":
                parsed.get(
                    "emd_amount"
                ),
            "bid_securing_declaration": (
                0
                if extracted_emd_required == 1
                else tender.bid_securing_declaration
            ),
            "emd_mode":
                tender.emd_mode,
            "emd_mode_required": (
                extracted_emd_required == 1
            ),
            "splitting_applicable":
                extracted_splitting,
            "splitting_ratio":
                parsed.get(
                    "splitting_ratio"
                ),
        },
    }


@frappe.whitelist()
def apply_reviewed_tender_extraction(
    tender_name,
    emd_mode=None,
):
    tender = frappe.get_doc(
        "PEPL Tender",
        tender_name,
    )

    tender.check_permission("write")

    if tender.docstatus != 0:
        frappe.throw(
            _(
                "Reviewed Tender extraction can be "
                "applied only while the Tender is Draft."
            )
        )

    parsed = (
        _get_reviewed_tender_parsed_data(
            tender
        )
    )

    applied = {}

    publication_date = parsed.get(
        "publication_date"
    )

    if publication_date:
        parsed_publication_date = (
            _parse_date(
                publication_date
            )
        )

        if not parsed_publication_date:
            frappe.throw(
                _(
                    "Extracted Publication Date could not "
                    "be converted to ERPNext date format: {0}"
                ).format(
                    publication_date
                )
            )

        tender.publication_date = (
            parsed_publication_date
        )

        applied[
            "publication_date"
        ] = parsed_publication_date

    bid_end = parsed.get(
        "bid_end"
    )

    if bid_end:
        parsed_bid_end = (
            _parse_datetime(
                bid_end
            )
        )

        if not parsed_bid_end:
            frappe.throw(
                _(
                    "Extracted Bid Submission Deadline could not "
                    "be converted to ERPNext datetime format: {0}"
                ).format(
                    bid_end
                )
            )

        tender.bid_submission_deadline = (
            parsed_bid_end
        )

        applied[
            "bid_submission_deadline"
        ] = parsed_bid_end

    bid_opening = parsed.get(
        "bid_opening"
    )

    if bid_opening:
        parsed_bid_opening = (
            _parse_datetime(
                bid_opening
            )
        )

        if not parsed_bid_opening:
            frappe.throw(
                _(
                    "Extracted Bid Opening Date could not "
                    "be converted to ERPNext datetime format: {0}"
                ).format(
                    bid_opening
                )
            )

        tender.bid_opening_date = (
            parsed_bid_opening
        )

        applied[
            "bid_opening_date"
        ] = parsed_bid_opening

    emd_required = _reviewed_yes_no(
        parsed.get(
            "emd_required"
        )
    )

    if emd_required is not None:
        tender.emd_required = (
            emd_required
        )

        applied[
            "emd_required"
        ] = emd_required

        # PEPL business rule:
        # EMD and Bid Securing Declaration are mutually exclusive.
        #
        # If the reviewed Tender explicitly requires EMD,
        # clear any earlier/default Bid Securing Declaration
        # before saving the ERP Tender.
        if emd_required == 1:
            if tender.bid_securing_declaration:
                tender.bid_securing_declaration = 0

                applied[
                    "bid_securing_declaration"
                ] = 0

            selected_emd_mode = (
                str(emd_mode or "").strip()
                or str(
                    tender.emd_mode or ""
                ).strip()
            )

            if not selected_emd_mode:
                frappe.throw(
                    _(
                        "EMD Mode is required because the reviewed "
                        "Tender extraction confirms that EMD is required."
                    )
                )

            emd_mode_field = (
                frappe.get_meta(
                    "PEPL Tender"
                ).get_field(
                    "emd_mode"
                )
            )

            allowed_modes = []

            if (
                emd_mode_field
                and emd_mode_field.options
            ):
                allowed_modes = [
                    value.strip()
                    for value in str(
                        emd_mode_field.options
                    ).splitlines()
                    if value.strip()
                ]

            if (
                allowed_modes
                and selected_emd_mode
                not in allowed_modes
            ):
                frappe.throw(
                    _(
                        "Invalid EMD Mode: {0}. "
                        "Please select one of the configured "
                        "PEPL Tender EMD modes."
                    ).format(
                        selected_emd_mode
                    )
                )

            tender.emd_mode = (
                selected_emd_mode
            )

            applied[
                "emd_mode"
            ] = selected_emd_mode

    emd_amount = parsed.get(
        "emd_amount"
    )

    if emd_amount not in {
        None,
        "",
    }:
        try:
            clean_amount = float(
                str(
                    emd_amount
                ).replace(
                    ",",
                    "",
                )
            )

            tender.emd_amount = (
                clean_amount
            )

            applied[
                "emd_amount"
            ] = clean_amount

        except ValueError:
            frappe.throw(
                _(
                    "Extracted EMD Amount is not numeric: {0}"
                ).format(
                    emd_amount
                )
            )

    splitting = _reviewed_yes_no(
        parsed.get(
            "splitting_applied"
        )
    )

    if splitting is not None:
        tender.splitting_applicable = (
            splitting
        )

        applied[
            "splitting_applicable"
        ] = splitting

    tender.save()

    tender.add_comment(
        "Info",
        (
            "Reviewed Tender extraction applied "
            "to deterministic ERP fields. "
            f"Applied values: {json.dumps(applied)}"
        ),
    )

    return {
        "applied": applied,
        "splitting_ratio":
            parsed.get(
                "splitting_ratio"
            ),
    }


@frappe.whitelist()
def generate_tender_word(
    tender_name,
):
    tender = frappe.get_doc(
        "PEPL Tender",
        tender_name,
    )

    tender.check_permission("read")

    if tender.tender_ingestion_status != "Reviewed":
        frappe.throw(
            _(
                "Complete Tender extraction review before "
                "generating the editable Word document."
            )
        )

    extraction = {}

    if tender.tender_extraction_json:
        try:
            extraction = json.loads(
                tender.tender_extraction_json
            )
        except json.JSONDecodeError:
            extraction = {}

    parsed = extraction.get(
        "parsed",
        {},
    )

    document = _load_word_letterhead()
    _set_docx_defaults(document)

    title = _add_docx_heading(
        document,
        "TENDER REVIEW / BID PREPARATION DOCUMENT",
        level=1,
        centered=True,
    )

    paragraph = document.add_paragraph()
    paragraph.add_run(
        "Generated automatically from uploaded Tender documents. "
    ).bold = True
    paragraph.add_run(
        "All extracted information must be reviewed against the "
        "original tender before bid submission."
    )

    _add_docx_heading(
        document,
        "1. Tender Identification",
        level=2,
    )

    _add_key_value_table(
        document,
        [
            (
                "PEPL Tender",
                tender.name,
            ),
            (
                "Tender / NIT Reference",
                tender.nit_number
                or parsed.get("bid_number"),
            ),
            (
                "Tender Title",
                tender.tender_title,
            ),
            (
                "Customer",
                tender.customer,
            ),
            (
                "Sector",
                tender.sector,
            ),
            (
                "Portal URL",
                tender.tender_portal_url,
            ),
        ],
    )

    _add_docx_heading(
        document,
        "2. Critical Dates",
        level=2,
    )

    _add_key_value_table(
        document,
        [
            (
                "Publication Date",
                tender.publication_date
                or parsed.get("publication_date"),
            ),
            (
                "Bid Submission Deadline",
                tender.bid_submission_deadline
                or parsed.get("bid_end"),
            ),
            (
                "Bid Opening",
                tender.bid_opening_date
                or parsed.get("bid_opening"),
            ),
            (
                "Pre-Bid Meeting",
                tender.pre_bid_meeting_date,
            ),
        ],
    )

    _add_docx_heading(
        document,
        "3. Extracted Organisation Information",
        level=2,
    )

    _add_key_value_table(
        document,
        [
            (
                "Ministry",
                parsed.get("ministry"),
            ),
            (
                "Department",
                parsed.get("department"),
            ),
            (
                "Organisation",
                parsed.get("organisation"),
            ),
            (
                "Total Quantity Found in Tender",
                parsed.get("total_quantity"),
            ),
        ],
    )

    _add_docx_heading(
        document,
        "4. Tender Items in ERPNext",
        level=2,
    )

    if tender.items:
        table = document.add_table(
            rows=1,
            cols=7,
        )

        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        headings = [
            "Item",
            "PL No",
            "Drawing",
            "Specification",
            "Quantity",
            "UOM",
            "Vendor Approval",
        ]

        for index, heading in enumerate(
            headings
        ):
            table.rows[0].cells[index].text = heading

        for item in tender.items:
            cells = table.add_row().cells

            values = [
                item.item,
                item.pl_no,
                item.drawing_no,
                item.current_specification,
                item.quantity,
                item.uom,
                (
                    f"{item.vendor_approval_stage or '-'} / "
                    f"{item.vendor_approval_health or '-'}"
                ),
            ]

            for index, value in enumerate(
                values
            ):
                cells[index].text = (
                    str(value)
                    if value not in (
                        None,
                        "",
                    )
                    else "-"
                )
    else:
        document.add_paragraph(
            "Tender Items have not yet been mapped to ERPNext Item Master."
        )

    _add_docx_heading(
        document,
        "5. EMD / Bid Security",
        level=2,
    )

    _add_key_value_table(
        document,
        [
            (
                "Bid Securing Declaration",
                "Yes"
                if tender.bid_securing_declaration
                else "No",
            ),
            (
                "ERP EMD Required",
                "Yes"
                if tender.emd_required
                else "No",
            ),
            (
                "ERP EMD Amount",
                tender.emd_amount,
            ),
            (
                "Extracted EMD Required",
                parsed.get("emd_required"),
            ),
            (
                "Extracted EMD Amount",
                parsed.get("emd_amount"),
            ),
            (
                "EMD Mode",
                tender.emd_mode,
            ),
            (
                "Extracted ePBG Percentage",
                parsed.get(
                    "epbg_percentage"
                ),
            ),
            (
                "Extracted ePBG Duration (Months)",
                parsed.get(
                    "epbg_duration_months"
                ),
            ),
            (
                "Extracted Splitting Applied",
                parsed.get(
                    "splitting_applied"
                ),
            ),
            (
                "Extracted Splitting Ratio",
                parsed.get(
                    "splitting_ratio"
                ),
            ),
            (
                "Extracted Bid Validity",
                parsed.get(
                    "bid_validity"
                ),
            ),
        ],
    )

    _add_docx_heading(
        document,
        "6. Required Bid Documents",
        level=2,
    )

    if tender.bid_documents:
        table = document.add_table(
            rows=1,
            cols=4,
        )

        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        for index, heading in enumerate(
            [
                "Document",
                "Source",
                "Mandatory",
                "Attached",
            ]
        ):
            table.rows[0].cells[index].text = heading

        for row in tender.bid_documents:
            cells = table.add_row().cells
            cells[0].text = (
                row.document_type
                or row.document_name
                or "-"
            )
            cells[1].text = row.document_source or "-"
            cells[2].text = (
                "Yes"
                if row.is_mandatory
                else "No"
            )
            cells[3].text = (
                "Yes"
                if row.is_attached
                else "No"
            )
    else:
        document.add_paragraph(
            "Bid-document checklist has not yet been generated."
        )

    _add_docx_heading(
        document,
        "7. Supporting Tender Documents",
        level=2,
    )

    _add_supporting_document_register(
        document,
        tender,
    )

    _add_docx_heading(
        document,
        "8. Supporting Document Content",
        level=2,
    )

    _add_supporting_document_excerpts(
        document,
        tender,
    )

    _add_docx_heading(
        document,
        "9. Review Exceptions",
        level=2,
    )

    _add_manual_review_register(
        document,
        tender,
    )

    _add_docx_heading(
        document,
        "10. Automatic Reading Warnings",
        level=2,
    )

    if tender.tender_ingestion_warnings:
        document.add_paragraph(
            tender.tender_ingestion_warnings
        )
    else:
        document.add_paragraph(
            "No automatic-reading warnings were recorded."
        )

    _add_docx_heading(
        document,
        "11. PEPL Manual Review / Sign-Off",
        level=2,
    )

    document.add_paragraph(
        "Tender reference verified: ______________________________"
    )
    document.add_paragraph(
        "Technical requirements verified: _________________________"
    )
    document.add_paragraph(
        "Commercial requirements verified: ________________________"
    )
    document.add_paragraph(
        "Bid documents verified: _________________________________"
    )
    document.add_paragraph(
        "Reviewed by: ______________________  Date: ______________"
    )

    output = BytesIO()
    document.save(output)

    output.seek(0)

    safe_reference = re.sub(
        r"[^A-Za-z0-9._-]+",
        "_",
        tender.nit_number or tender.name,
    )

    filename = (
        f"PEPL_Tender_Review_{safe_reference}.docx"
    )

    file_url = _attach_content(
        content=output.getvalue(),
        filename=filename,
        tender_name=tender.name,
        attached_to_field="generated_tender_word",
    )

    tender.db_set(
        "generated_tender_word",
        file_url,
        update_modified=True,
    )

    tender.add_comment(
        "Comment",
        text=(
            "Editable Tender Word document generated by "
            f"{frappe.session.user}: {file_url}"
        ),
    )

    return {
        "file_url": file_url,
        "filename": filename,
    }
